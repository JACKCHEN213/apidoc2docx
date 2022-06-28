# -*- coding: utf-8 -*-
import json
import re
from pydantic import BaseModel, Field, Extra
from typing import List
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX


class Generator(BaseModel):
    name: str = Field(default='apidoc', description='生成文档工具名称')
    time: str = Field(default='', description='文档生成时间')
    version: str = Field(default='1.0.0', description='生成工具版本')
    url: str = Field(default='', description='生成工具的官方文档')


class Project(BaseModel):
    name: str = Field(..., description='项目的名称')
    version: str = Field(..., description='项目的版本')
    description: str = Field(default='', description='项目的描述')
    url: str = Field(..., description='项目的路由前缀')
    generator: Generator = Field(..., description='文档生成器')

    def write_to_docx(self, document: Document):
        document.add_heading(f'{self.name} v{self.version}', 1)
        document.add_paragraph(self.description)

    def excel_name(self, output: str):
        return f'{output}/{self.name}_V{self.version}.docx'


class BaseItem(BaseModel):
    group: str = Field(default='Parameter', description='参数分组')
    type: str = Field(default='String', description='参数类型')
    optional: bool = Field(default=True, description='参数是否必填')
    field: str = Field(default='', description='参数名称')
    description: str = Field(default='', description='参数的描述')


class Fields(BaseModel):
    Parameter: List[BaseItem] = Field(default=list(), description='参数内容')

    class Config:
        extra = Extra.allow


class Examples(BaseModel):
    title: str = Field(default='', description='示例标题')
    content: str = Field(default='', description='示例内容')
    type: str = Field(default='json', description='示例内容的类别')


class Parameter(BaseModel):
    fields: Fields = Field(default=BaseModel(), description='请求参数字段')
    examples: List[Examples] = Field(default=list(), description='请求示例')


class Success(BaseModel):
    fields: Fields = Field(default=BaseModel(), description='响应参数字段')
    examples: List[Examples] = Field(default=list(), description='响应示例')


class Item(BaseModel):
    type: str = Field(..., description='接口请求方式')
    url: str = Field(..., description='接口请求路径')
    title: str = Field(..., description='接口标题')
    description: str = Field(default='', description='接口描述')
    group: str = Field(..., description='接口分组')
    name: str = Field(..., description='接口名称')
    parameter: Parameter = Field(default=BaseModel(), description='接口请求参数')
    success: Success = Field(default=BaseModel(), description='接口响应参数')
    version: str = Field(default='0.0.0', description='接口版本')
    filename: str = Field(default='', description='接口文件路径')
    groupTitle: str = Field(default='', description='接口分组标题')


class Group(BaseModel):
    name: str = Field(..., description='接口分组名称')
    items: List[Item] = Field(default=list(), description='接口列表')

    def write_to_docs(self, document: Document, project: Project):
        document.add_heading(self.name, 1)
        for item in self.items:
            document.add_heading(item.title, 2)
            p = document.add_paragraph(re.sub('<[^>]+?>', '', item.description).strip())
            p.paragraph_format.space_after = Pt(0)
            p = document.add_paragraph()
            run = p.add_run(' ' + item.type.upper() + ' ')
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = '宋体'
            run.font.bold = True
            run.font.highlight_color = WD_COLOR_INDEX.TEAL
            p.paragraph_format.line_spacing = Pt(18)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.space_before = Pt(5)

            # 添加请求url
            cell = document.add_table(rows=1, cols=1, style='Table Grid').cell(0, 0)
            cell.text = project.url + item.url

            parameter = item.parameter.dict()
            self._write_parameter(document, parameter)
            success = item.success.dict()
            self._write_success(document, success)

    def _write_parameter(self, document: Document, parameter: dict):
        # 请求参数
        if 'fields' in parameter:
            for param, attributes in parameter['fields'].items():
                if not attributes:
                    continue
                document.add_heading('参数-' + param, 3)
                table = document.add_table(cols=4, rows=(1 + len(attributes)), style='Table Grid')
                table.cell(0, 0).text = '参数名称'
                table.cell(0, 1).text = '是否必填'
                table.cell(0, 2).text = '参数类型'
                table.cell(0, 3).text = '参数说明'
                # 参数详情
                for row in range(len(attributes)):
                    attribute = attributes[row]
                    table.cell(row + 1, 0).text = attribute['field']
                    table.cell(row + 1, 1).text = '√' if not attribute['optional'] else ''
                    table.cell(row + 1, 2).text = attribute['type']
                    table.cell(row + 1, 3).text = re.sub('<[^>]+?>', '', attribute['description']).strip()
                self._set_table_style(table, len(attributes) + 1, 4)
        # 请求示例
        if 'examples' in parameter:
            document.add_heading('参数示例', 3)
            for example in parameter['examples']:
                document.add_paragraph(example['type'] + '-' + example['title'])
                cell = document.add_table(rows=1, cols=1, style='Table Grid').cell(0, 0)
                if example['type'] == 'json':
                    cell.text = json.dumps(json.loads(example['content']), indent=2, ensure_ascii=False)
                else:
                    cell.text = example['content']
                cell.paragraphs[0].runs[0].font.name = 'arial'

    def _write_success(self, document: Document, success: dict):
        # 响应参数
        document.add_heading('响应参数', 3)
        if 'fields' in success:
            for param, attributes in success['fields'].items():
                if not attributes:
                    continue
                h = document.add_heading('响应参数-' + param, 4)
                h.paragraph_format.space_before = Pt(0)
                table = document.add_table(cols=4, rows=(1 + len(attributes)), style='Table Grid')
                table.cell(0, 0).text = '参数名称'
                table.cell(0, 1).text = '是否必返'
                table.cell(0, 2).text = '参数类型'
                table.cell(0, 3).text = '参数说明'
                # 参数详情
                for row in range(len(attributes)):
                    attribute = attributes[row]
                    table.cell(row + 1, 0).text = attribute['field']
                    table.cell(row + 1, 1).text = '√' if not attribute['optional'] else ''
                    table.cell(row + 1, 2).text = attribute['type']
                    table.cell(row + 1, 3).text = re.sub('<[^>]+?>', '', attribute['description']).strip()
                self._set_table_style(table, len(attributes) + 1, 4)
        # 响应示例
        if 'examples' in success:
            document.add_heading('响应示例', 3)
            for example in success['examples']:
                h = document.add_heading('响应示例' + '-' + example['title'], 4)
                h.paragraph_format.space_before = Pt(0)
                cell = document.add_table(rows=1, cols=1, style='Table Grid').cell(0, 0)
                if example['type'] == 'json':
                    cell.text = json.dumps(json.loads(example['content']), indent=2, ensure_ascii=False)
                else:
                    cell.text = example['content']
                cell.paragraphs[0].runs[0].font.name = 'arial'
    
    @staticmethod
    def _set_table_style(table, rows: int, cols: int):
        table.cell(0, 0).width = Inches(2)
        table.cell(0, 1).width = Inches(1)
        table.cell(0, 2).width = Inches(2)
        table.cell(0, 3).width = Inches(5)
        for row in range(rows):
            for col in range(cols):
                table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
